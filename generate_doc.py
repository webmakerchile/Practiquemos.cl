from docx import Document
from docx.shared import Inches, Pt, Cm, RGBColor
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.enum.table import WD_TABLE_ALIGNMENT
from docx.enum.section import WD_ORIENT
from docx.oxml.ns import qn
from docx.oxml import OxmlElement
import os

doc = Document()

style = doc.styles['Normal']
font = style.font
font.name = 'Calibri'
font.size = Pt(11)
font.color.rgb = RGBColor(0x33, 0x33, 0x33)
style.paragraph_format.space_after = Pt(6)
style.paragraph_format.line_spacing = 1.15

for section in doc.sections:
    section.top_margin = Cm(2)
    section.bottom_margin = Cm(2)
    section.left_margin = Cm(2.5)
    section.right_margin = Cm(2.5)

GREEN = RGBColor(0x1B, 0x8A, 0x2F)
DARK_GREEN = RGBColor(0x14, 0x6B, 0x23)
ORANGE = RGBColor(0xF5, 0x9E, 0x0B)
DARK_GRAY = RGBColor(0x33, 0x33, 0x33)
MEDIUM_GRAY = RGBColor(0x66, 0x66, 0x66)
BLUE = RGBColor(0x1D, 0x4E, 0xD8)
RED = RGBColor(0xDC, 0x26, 0x26)
WHITE = RGBColor(0xFF, 0xFF, 0xFF)

def add_colored_heading(text, level=1, color=GREEN):
    heading = doc.add_heading(text, level=level)
    for run in heading.runs:
        run.font.color.rgb = color
        run.font.name = 'Calibri'
    return heading

def add_separator():
    p = doc.add_paragraph()
    p.paragraph_format.space_before = Pt(4)
    p.paragraph_format.space_after = Pt(4)
    run = p.add_run('_' * 80)
    run.font.color.rgb = RGBColor(0xCC, 0xCC, 0xCC)
    run.font.size = Pt(6)

def set_cell_shading(cell, color_hex):
    shading = OxmlElement('w:shd')
    shading.set(qn('w:fill'), color_hex)
    shading.set(qn('w:val'), 'clear')
    cell._tc.get_or_add_tcPr().append(shading)

def style_table_cell(cell, text, bold=False, color=DARK_GRAY, bg_color=None, align=WD_ALIGN_PARAGRAPH.LEFT, font_size=10):
    cell.text = ''
    p = cell.paragraphs[0]
    p.alignment = align
    run = p.add_run(text)
    run.font.name = 'Calibri'
    run.font.size = Pt(font_size)
    run.font.color.rgb = color
    run.bold = bold
    if bg_color:
        set_cell_shading(cell, bg_color)

logo_path = 'attached_assets/Logo_sin_fondo_1773872982719.png'

p_logo = doc.add_paragraph()
p_logo.alignment = WD_ALIGN_PARAGRAPH.CENTER
if os.path.exists(logo_path):
    run_logo = p_logo.add_run()
    run_logo.add_picture(logo_path, width=Inches(2.2))

doc.add_paragraph()

p_title = doc.add_paragraph()
p_title.alignment = WD_ALIGN_PARAGRAPH.CENTER
run_t = p_title.add_run('DOCUMENTO DE ENTREGA')
run_t.font.size = Pt(28)
run_t.font.color.rgb = GREEN
run_t.bold = True
run_t.font.name = 'Calibri'

p_sub = doc.add_paragraph()
p_sub.alignment = WD_ALIGN_PARAGRAPH.CENTER
run_s = p_sub.add_run('Practiquemos.cl')
run_s.font.size = Pt(22)
run_s.font.color.rgb = DARK_GREEN
run_s.bold = True
run_s.font.name = 'Calibri'

p_desc = doc.add_paragraph()
p_desc.alignment = WD_ALIGN_PARAGRAPH.CENTER
run_d = p_desc.add_run('Plataforma de preparación para el examen de licencia de conducir en Chile')
run_d.font.size = Pt(12)
run_d.font.color.rgb = MEDIUM_GRAY
run_d.italic = True
run_d.font.name = 'Calibri'

doc.add_paragraph()
add_separator()

p_info = doc.add_paragraph()
p_info.alignment = WD_ALIGN_PARAGRAPH.CENTER
p_info.paragraph_format.space_before = Pt(12)

run_date = p_info.add_run('Fecha de entrega: 18 de marzo de 2026\n')
run_date.font.size = Pt(11)
run_date.font.color.rgb = MEDIUM_GRAY
run_date.font.name = 'Calibri'

run_ver = p_info.add_run('Versión: 1.0\n')
run_ver.font.size = Pt(11)
run_ver.font.color.rgb = MEDIUM_GRAY
run_ver.font.name = 'Calibri'

run_dev = p_info.add_run('Desarrollado por: WebMakerChile')
run_dev.font.size = Pt(11)
run_dev.font.color.rgb = GREEN
run_dev.bold = True
run_dev.font.name = 'Calibri'

doc.add_page_break()

# ===================== ÍNDICE =====================
add_colored_heading('Índice', level=1)
add_separator()

indice_items = [
    '1. Bienvenida',
    '2. Cómo empezar',
    '3. Funcionalidades de la plataforma',
    '4. Panel de administrador',
    '5. Gestión de preguntas',
    '6. Planes y pagos',
    '7. Configuración de exámenes por licencia',
    '8. Seguridad y protección de datos',
    '9. Estado actual y pendientes',
    '10. Contacto y soporte',
]

for item in indice_items:
    p = doc.add_paragraph(item)
    p.paragraph_format.space_after = Pt(4)
    for run in p.runs:
        run.font.color.rgb = DARK_GRAY
        run.font.size = Pt(11)

doc.add_page_break()

# ===================== 1. BIENVENIDA =====================
add_colored_heading('1. Bienvenida', level=1)
add_separator()

p = doc.add_paragraph()
run = p.add_run('Estimado cliente,')
run.font.size = Pt(11)
run.font.color.rgb = DARK_GRAY

p = doc.add_paragraph(
    'Es un placer hacer entrega de su plataforma Practiquemos.cl, una aplicación móvil y web '
    'diseñada para ayudar a las personas a prepararse para el examen de licencia de conducir en Chile.'
)

p = doc.add_paragraph(
    'La plataforma ha sido desarrollada con las más modernas tecnologías y está lista para funcionar '
    'en navegadores web, dispositivos iOS (iPhone/iPad) y Android. A continuación, encontrará toda la '
    'información necesaria para comenzar a utilizar y administrar su plataforma.'
)

doc.add_paragraph()

# ===================== 2. CÓMO EMPEZAR =====================
add_colored_heading('2. Cómo empezar', level=1)
add_separator()

add_colored_heading('Paso 1: Crear su cuenta', level=2, color=DARK_GREEN)
p = doc.add_paragraph()
p.add_run('Para acceder como administrador de la plataforma, primero debe crear una cuenta de usuario:').font.color.rgb = DARK_GRAY

steps = [
    'Ingrese a la aplicación desde su navegador o dispositivo móvil.',
    'Toque el botón "Registrarse".',
    'Complete sus datos personales: nombre de usuario, contraseña, nombre completo y correo electrónico.',
    'Una vez registrado, comunique su nombre de usuario al equipo de WebMakerChile.',
]
for i, step in enumerate(steps, 1):
    p = doc.add_paragraph(f'{i}. {step}')
    p.paragraph_format.left_indent = Cm(1)

add_colored_heading('Paso 2: Activación de permisos de administrador', level=2, color=DARK_GREEN)
p = doc.add_paragraph(
    'Una vez que WebMakerChile reciba su nombre de usuario, le otorgará los permisos de administrador '
    'desde el servidor. A partir de ese momento, tendrá acceso completo al Panel de Administrador '
    'desde el menú principal de la aplicación.'
)

p = doc.add_paragraph()
run = p.add_run('Importante: ')
run.bold = True
run.font.color.rgb = RED
run = p.add_run(
    'No comparta sus credenciales de acceso con terceros. Usted es responsable de la seguridad '
    'de su cuenta de administrador.'
)

doc.add_paragraph()

# ===================== 3. FUNCIONALIDADES =====================
add_colored_heading('3. Funcionalidades de la plataforma', level=1)
add_separator()

add_colored_heading('Para los usuarios', level=2, color=DARK_GREEN)

features = [
    ('6 tipos de licencia', 'Clase B, A2, A4, C, D y E, cada una con preguntas específicas para su categoría.'),
    ('1.591 preguntas', 'Base de datos completa que incluye 306 preguntas oficiales de CONASET.'),
    ('Exámenes de práctica', 'Dos modalidades: Simulación (con temporizador, simula el examen real) y Aprendizaje (con explicaciones detalladas después de cada respuesta).'),
    ('Temario completo', '11 capítulos de material de estudio con secciones detalladas que cubren todos los temas del examen.'),
    ('Historial de exámenes', 'Los usuarios pueden consultar todos sus resultados anteriores y ver su evolución.'),
    ('Preguntas favoritas', 'Función para guardar preguntas y repasarlas en cualquier momento.'),
    ('Progreso por categoría', 'Seguimiento visual del avance en cada tema del examen.'),
    ('Lectura en voz alta', 'Las preguntas y explicaciones se pueden escuchar con voz natural de alta calidad.'),
    ('Sonidos interactivos', 'Efectos de sonido al responder correcta o incorrectamente para una experiencia más dinámica.'),
    ('Mascota copiloto', 'Un compañero animado que acompaña y motiva al usuario durante el estudio.'),
    ('Imágenes ilustrativas', '68 imágenes en las preguntas para facilitar la comprensión visual de las situaciones de tránsito.'),
    ('Barajado de opciones', 'Las alternativas se mezclan aleatoriamente en cada examen para evitar memorización por posición.'),
]

for title, desc in features:
    p = doc.add_paragraph()
    run_title = p.add_run(f'• {title}: ')
    run_title.bold = True
    run_title.font.color.rgb = DARK_GREEN
    run_desc = p.add_run(desc)
    run_desc.font.color.rgb = DARK_GRAY

doc.add_paragraph()

# ===================== 4. PANEL DE ADMINISTRADOR =====================
add_colored_heading('4. Panel de administrador', level=1)
add_separator()

p = doc.add_paragraph(
    'Como administrador, usted tiene control total sobre la plataforma. El panel de administración '
    'le permite gestionar usuarios, preguntas y contenido de manera intuitiva.'
)

add_colored_heading('Gestión de usuarios', level=2, color=DARK_GREEN)

user_features = [
    ('Crear usuarios', 'Registre nuevos usuarios con nombre, correo electrónico, contraseña, tipo de licencia, plan y rol.'),
    ('Editar usuarios', 'Modifique el nombre, correo electrónico, contraseña, plan (gratuito o premium) y rol de cualquier usuario.'),
    ('Eliminar usuarios', 'Elimine usuarios con una confirmación de seguridad para evitar borrados accidentales.'),
    ('Buscar usuarios', 'Encuentre usuarios rápidamente por nombre, nombre de usuario o correo electrónico.'),
    ('Filtrar usuarios', 'Filtre la lista por categoría: todos, premium, gratuitos o administradores.'),
    ('Ver detalles', 'Consulte la información completa de cada usuario: fecha de registro, último inicio de sesión, vigencia del plan y tipo de licencia.'),
    ('Otorgar premium', 'Cambie manualmente el plan de cualquier usuario de gratuito a premium (10 o 30 días) directamente desde el panel.'),
]

for title, desc in user_features:
    p = doc.add_paragraph()
    run_title = p.add_run(f'• {title}: ')
    run_title.bold = True
    run_title.font.color.rgb = DARK_GREEN
    run_desc = p.add_run(desc)

add_colored_heading('Estadísticas en tiempo real', level=2, color=DARK_GREEN)
p = doc.add_paragraph(
    'El panel muestra tarjetas interactivas con estadísticas actualizadas: total de usuarios, '
    'usuarios premium activos, usuarios gratuitos y cantidad de administradores. Al tocar cada '
    'tarjeta, se filtra automáticamente la lista de usuarios correspondiente.'
)

doc.add_paragraph()

# ===================== 5. GESTIÓN DE PREGUNTAS =====================
add_colored_heading('5. Gestión de preguntas', level=1)
add_separator()

p = doc.add_paragraph(
    'Desde el panel de administrador puede acceder a la gestión completa de las 1.591 preguntas de la plataforma:'
)

question_features = [
    ('Buscar y filtrar', 'Encuentre preguntas por texto, categoría, tipo de licencia, dificultad y origen (oficial o no oficial).'),
    ('Crear preguntas', 'Agregue nuevas preguntas con sus opciones de respuesta, respuesta correcta, explicación y categoría.'),
    ('Editar preguntas', 'Modifique cualquier aspecto de las preguntas existentes.'),
    ('Desactivar preguntas', 'Las preguntas no se eliminan, sino que se desactivan (se ocultan) para que puedan reactivarse en el futuro.'),
    ('Preguntas oficiales', 'Las preguntas provenientes de CONASET están marcadas como oficiales para diferenciarlas del contenido propio.'),
]

for title, desc in question_features:
    p = doc.add_paragraph()
    run_title = p.add_run(f'• {title}: ')
    run_title.bold = True
    run_title.font.color.rgb = DARK_GREEN
    run_desc = p.add_run(desc)

doc.add_paragraph()

# ===================== 6. PLANES Y PAGOS =====================
add_colored_heading('6. Planes y pagos', level=1)
add_separator()

add_colored_heading('Modelo de negocio', level=2, color=DARK_GREEN)
p = doc.add_paragraph(
    'La plataforma funciona con un modelo freemium: los usuarios pueden acceder de forma gratuita '
    'con funcionalidad limitada, y desbloquear el contenido completo con un plan premium.'
)

table = doc.add_table(rows=4, cols=3)
table.alignment = WD_TABLE_ALIGNMENT.CENTER
table.style = 'Table Grid'

headers = ['Plan', 'Precio', 'Duración']
for i, header in enumerate(headers):
    style_table_cell(table.rows[0].cells[i], header, bold=True, color=WHITE, bg_color='1B8A2F', align=WD_ALIGN_PARAGRAPH.CENTER)

plans_data = [
    ('Gratuito', 'Sin costo', 'Permanente (acceso limitado)'),
    ('Premium 10 días', '$2.990 CLP', '10 días de acceso completo'),
    ('Premium 30 días', '$4.990 CLP', '30 días de acceso completo'),
]

for row_idx, (plan, price, dur) in enumerate(plans_data, 1):
    style_table_cell(table.rows[row_idx].cells[0], plan, bold=True, align=WD_ALIGN_PARAGRAPH.CENTER)
    style_table_cell(table.rows[row_idx].cells[1], price, align=WD_ALIGN_PARAGRAPH.CENTER)
    style_table_cell(table.rows[row_idx].cells[2], dur, align=WD_ALIGN_PARAGRAPH.CENTER)

doc.add_paragraph()

add_colored_heading('Métodos de pago', level=2, color=DARK_GREEN)

p = doc.add_paragraph()
run = p.add_run('Web y Android: ')
run.bold = True
run.font.color.rgb = DARK_GREEN
p.add_run('Los pagos se procesan a través de Mercado Pago (Checkout Pro). Los usuarios pueden pagar directamente desde la aplicación y su plan se activa de forma automática.')

p = doc.add_paragraph()
run = p.add_run('iOS (App Store): ')
run.bold = True
run.font.color.rgb = DARK_GREEN
p.add_run('Para compras dentro de la aplicación en iPhone y iPad, se utiliza RevenueCat integrado con el sistema de compras de Apple. Esta configuración es gestionada por WebMakerChile.')

p = doc.add_paragraph()
run = p.add_run('Asignación manual: ')
run.bold = True
run.font.color.rgb = DARK_GREEN
p.add_run('Como administrador, usted puede otorgar el plan premium directamente a cualquier usuario desde el panel de administración, sin que el usuario tenga que realizar un pago.')

doc.add_paragraph()

# ===================== 7. CONFIGURACIÓN DE EXÁMENES =====================
add_colored_heading('7. Configuración de exámenes por licencia', level=1)
add_separator()

p = doc.add_paragraph(
    'Cada tipo de licencia tiene su propia configuración de examen, basada en los parámetros oficiales:'
)

table2 = doc.add_table(rows=7, cols=4)
table2.alignment = WD_TABLE_ALIGNMENT.CENTER
table2.style = 'Table Grid'

exam_headers = ['Licencia', 'Preguntas', 'Puntaje para aprobar', 'Tiempo']
for i, header in enumerate(exam_headers):
    style_table_cell(table2.rows[0].cells[i], header, bold=True, color=WHITE, bg_color='1B8A2F', align=WD_ALIGN_PARAGRAPH.CENTER)

exam_data = [
    ('Clase B', '35', '33/38 (87%)', '45 minutos'),
    ('Clase A2', '20', '16/20 (80%)', '30 minutos'),
    ('Clase A4', '35', '70%', '45 minutos'),
    ('Clase C', '20', '15/20 (75%)', '30 minutos'),
    ('Clase D', '12', '9/12 (75%)', '20 minutos'),
    ('Clase E', '10', '7/10 (70%)', '20 minutos'),
]

for row_idx, (lic, qs, score, time) in enumerate(exam_data, 1):
    style_table_cell(table2.rows[row_idx].cells[0], lic, bold=True, align=WD_ALIGN_PARAGRAPH.CENTER)
    style_table_cell(table2.rows[row_idx].cells[1], qs, align=WD_ALIGN_PARAGRAPH.CENTER)
    style_table_cell(table2.rows[row_idx].cells[2], score, align=WD_ALIGN_PARAGRAPH.CENTER)
    style_table_cell(table2.rows[row_idx].cells[3], time, align=WD_ALIGN_PARAGRAPH.CENTER)

p = doc.add_paragraph()
p.paragraph_format.space_before = Pt(8)
run = p.add_run('Nota: ')
run.bold = True
run.font.color.rgb = ORANGE
p.add_run('En la Clase B, 3 preguntas valen 2 puntos cada una, por lo que el puntaje máximo es 38 puntos (no 35).')

doc.add_paragraph()

# ===================== 8. SEGURIDAD =====================
add_colored_heading('8. Seguridad y protección de datos', level=1)
add_separator()

security_items = [
    ('Contraseñas encriptadas', 'Todas las contraseñas se almacenan de forma segura utilizando encriptación bcrypt. Nunca se guardan en texto plano.'),
    ('Sesiones seguras', 'Las sesiones de usuario se gestionan mediante tokens de autenticación seguros.'),
    ('Protección de datos', 'La base de datos PostgreSQL almacena toda la información de forma segura con respaldos automáticos.'),
    ('Política de privacidad', 'La aplicación incluye una política de privacidad y términos de servicio accesibles desde la sección Legal.'),
    ('Roles de usuario', 'El sistema cuenta con tres niveles de acceso: usuario estándar, administrador y super administrador (uso exclusivo de WebMakerChile).'),
]

for title, desc in security_items:
    p = doc.add_paragraph()
    run_title = p.add_run(f'• {title}: ')
    run_title.bold = True
    run_title.font.color.rgb = DARK_GREEN
    run_desc = p.add_run(desc)

doc.add_paragraph()

# ===================== 9. ESTADO ACTUAL Y PENDIENTES =====================
add_colored_heading('9. Estado actual y pendientes', level=1)
add_separator()

add_colored_heading('Listo y funcionando', level=2, color=DARK_GREEN)

ready_items = [
    'Aplicación web completamente funcional y publicada.',
    'Sistema de registro e inicio de sesión de usuarios.',
    'Base de datos con 1.591 preguntas (incluidas 306 oficiales de CONASET).',
    'Exámenes de práctica con modo simulación y modo aprendizaje.',
    'Temario completo con 11 capítulos de estudio.',
    'Historial de exámenes, favoritos y progreso por categoría.',
    'Panel de administrador con gestión completa de usuarios y preguntas.',
    'Sistema de pagos con Mercado Pago (web y Android).',
    'Integración con RevenueCat para compras en iOS (en proceso de configuración).',
    'Lectura en voz alta con inteligencia artificial (voz Nova de OpenAI).',
    'Mascota copiloto con animaciones interactivas.',
    'Política de privacidad y términos de servicio.',
]

for item in ready_items:
    p = doc.add_paragraph(f'✓  {item}')
    p.runs[0].font.color.rgb = DARK_GREEN

add_colored_heading('Pendiente (configuración externa)', level=2, color=ORANGE)

pending_items = [
    ('Publicación en App Store (iOS)', 'Requiere una cuenta de Apple Developer ($99 USD/año). WebMakerChile se encargará de la compilación y envío a Apple.'),
    ('Publicación en Google Play Store (Android)', 'Requiere una cuenta de Google Play Developer ($25 USD, pago único). WebMakerChile se encargará de la compilación y publicación.'),
    ('Configuración de RevenueCat con App Store Connect', 'Necesario para que las compras dentro de la aplicación funcionen en iOS. Requiere la cuenta de Apple Developer activa.'),
]

for title, desc in pending_items:
    p = doc.add_paragraph()
    run_title = p.add_run(f'⚠  {title}: ')
    run_title.bold = True
    run_title.font.color.rgb = ORANGE
    run_desc = p.add_run(desc)
    run_desc.font.color.rgb = DARK_GRAY

doc.add_paragraph()

# ===================== 10. CONTACTO =====================
add_colored_heading('10. Contacto y soporte', level=1)
add_separator()

p = doc.add_paragraph(
    'Para cualquier consulta técnica, soporte o solicitudes relacionadas con la plataforma, '
    'puede comunicarse a través de los siguientes canales:'
)

p = doc.add_paragraph()
run = p.add_run('Correo electrónico: ')
run.bold = True
run.font.color.rgb = DARK_GREEN
p.add_run('practiquemos.cl@gmail.com')

p = doc.add_paragraph()
run = p.add_run('Desarrollador: ')
run.bold = True
run.font.color.rgb = DARK_GREEN
p.add_run('WebMakerChile')

doc.add_paragraph()
doc.add_paragraph()
add_separator()

p_footer = doc.add_paragraph()
p_footer.alignment = WD_ALIGN_PARAGRAPH.CENTER
p_footer.paragraph_format.space_before = Pt(20)

if os.path.exists(logo_path):
    run_logo2 = p_footer.add_run()
    run_logo2.add_picture(logo_path, width=Inches(1.2))

p_copy = doc.add_paragraph()
p_copy.alignment = WD_ALIGN_PARAGRAPH.CENTER
run_c = p_copy.add_run('© 2026 Practiquemos.cl — Todos los derechos reservados')
run_c.font.size = Pt(9)
run_c.font.color.rgb = MEDIUM_GRAY

p_dev = doc.add_paragraph()
p_dev.alignment = WD_ALIGN_PARAGRAPH.CENTER
run_dv = p_dev.add_run('Desarrollado con dedicación por WebMakerChile')
run_dv.font.size = Pt(9)
run_dv.font.color.rgb = GREEN
run_dv.italic = True

doc.save('Documento_Entrega_Practiquemos.docx')
print('Documento Word generado exitosamente.')
